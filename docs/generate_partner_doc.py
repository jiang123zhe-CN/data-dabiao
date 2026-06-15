# -*- coding: utf-8 -*-
"""生成给合伙人的非技术版平台说明文档（Word）——数据资产管理平台"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ===== 全局样式 =====
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# 标题样式
for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.name = '微软雅黑'
    heading_style.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

# ===== 封面 =====
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('数据资产管理平台')
run.font.size = Pt(28)
run.font.bold = True
run.font.name = '微软雅黑'
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('测试版操作指南')
run.font.size = Pt(16)
run.font.name = '微软雅黑'
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('版本：v1.0 测试版\n日期：2026年6月\n适用对象：金融机构数据治理负责人、合规负责人').font.size = Pt(11)

doc.add_page_break()

# ===== 一、平台是什么 =====
doc.add_heading('一、平台是什么', level=1)
doc.add_paragraph(
    '简单说：一个帮助金融机构自动完成数据资产"分类分级"的工具。\n\n'
    '金融机构（银行、保险、券商等）在日常经营中积累了海量数据——客户姓名、身份证号、'
    '交易流水、风险模型参数……这些数据的安全敏感程度各不相同。根据监管要求（'
    '《数据安全法》《个人信息保护法》《金融数据安全 数据安全分级指南》），'
    '机构必须对每一类数据资产进行梳理、分类、定级，并形成资产目录。\n\n'
    '过去这项工作主要靠人工：数据管理员打开数据库，逐表逐字段地判断"这个字段属于哪个业务域？'
    '敏感级别是L2还是L3？应该映射到哪个资产目录？"。一个中型银行可能有上万字段，'
    '纯人工做一遍需要数月，且容易出错。\n\n'
    '这个平台利用 AI（通义千问大模型）自动完成字段分类和目录映射，人工只需复核确认，'
    '效率可提升 5-10 倍。同时内置金融行业分类标准（8大分类体系）和 L1-L4 安全分级规则，'
    '确保分类分级结果符合监管要求。'
)

# ===== 二、为什么需要这个平台 =====
doc.add_heading('二、现在的痛点（为什么做这个）', level=1)

doc.add_paragraph('金融机构在数据资产分类分级上面临三个核心问题：')

problems = [
    ('工作量大、效率低',
     '一个中型金融机构的数据仓库里有数百张表、数千甚至上万个字段。数据管理员需要逐一判断每个字段的'
     '业务分类、安全级别、所属资产目录。纯人工方式耗时数月，且高度依赖个人经验。'),
    ('标准不统一',
     '不同部门、不同人对同一个字段的分类定级可能不同——"手机号"到底是L2还是L3？"交易金额"算交易域还是财务域？'
     '缺乏统一的规则引擎和标准参照，导致分类结果不一致，监管检查时难以自证合规。'),
    ('AI 结果不可信',
     '直接用大模型做分类虽然快，但模型可能"幻觉"——把身份证号误判为普通文本、把高风险字段标为低风险。'
     '监管要求数据分级必须有"人"的确认，纯 AI 输出无法直接作为合规依据。'),
]
for t, desc in problems:
    p = doc.add_paragraph()
    run = p.add_run(f'▸ {t}：')
    run.bold = True
    p.add_run(desc)

# ===== 三、平台怎么玩 =====
doc.add_heading('三、平台运作机制（五个角色）', level=1)

doc.add_paragraph('平台上有五类用户，权限逐级递增：')

doc.add_heading('角色一：数据录入员（data_entry）', level=2)
doc.add_paragraph(
    '负责将数据字段录入系统，是数据资产管理的"入口"。可以：\n'
    '① 手动录入字段信息（字段名、数据类型、所属表、数据库等）\n'
    '② 通过 Excel 批量导入字段\n'
    '③ 编辑自己录入的字段\n'
    '④ 查看字段的异常标记（系统自动检测未映射或信息不完整的字段）'
)

doc.add_heading('角色二：复核员（reviewer）', level=2)
doc.add_paragraph(
    '负责审核 AI 的建议和系统发现的异常，是"人机协同"中的"人"：\n'
    '① 查看系统自动检测的异常字段（未映射、信息缺失）\n'
    '② 审核 AI 自动映射建议——可以批准或拒绝\n'
    '③ 查看操作日志\n\n'
    '关键设计：AI 的映射建议不会直接生效，必须经过复核员批准才会写入系统。'
    '这确保了"AI 提建议，人做决策"的合规流程。'
)

doc.add_heading('角色三：数据管理员（data_admin）', level=2)
doc.add_paragraph(
    '负责资产目录和映射关系的日常管理：\n'
    '① 创建和维护资产目录树（如：客户域 → 客户基本信息 → 姓名、证件号……）\n'
    '② 手动建立字段与目录的映射关系\n'
    '③ 触发 AI 自动映射——系统会将未映射的字段与目录树一起发给大模型，'
    '由模型推荐每个字段应该归属到哪个目录节点\n'
    '④ 管理数据标签（打标）'
)

doc.add_heading('角色四：系统管理员（system_admin）', level=2)
doc.add_paragraph(
    '负责分类分级标准和规则的管理：\n'
    '① 管理金融数据分类标准（内置8大分类体系，可自定义扩展）\n'
    '② 管理 L1-L4 安全分级规则（关键字匹配、正则表达式匹配）\n'
    '③ 配置数据源连接信息'
)

doc.add_heading('角色五：超级管理员（admin）', level=2)
doc.add_paragraph(
    '拥有全部权限，额外负责：\n'
    '① 用户管理（创建、禁用、重置密码）\n'
    '② 查看报表和统计分析\n'
    '③ 合规审计'
)

doc.add_heading('一条字段的完整旅程', level=2)
doc.add_paragraph(
    '举个例子，方便理解整个流程：\n\n'
    '1️⃣ 数据录入员登录，导入客户信息表（dim_customer）的字段列表，包含"身份证号""手机号""邮箱"等字段\n\n'
    '2️⃣ 系统自动检测：这3个字段目前都没有映射到任何资产目录 → 标记为"异常字段"\n\n'
    '3️⃣ 数据管理员登录，点击"AI 自动映射"：系统将3个字段 + 目录树发给大模型，模型返回建议：\n'
    '    身份证号 → 客户域/客户标识信息（L4 机密）\n'
    '    手机号   → 客户域/客户联系信息（L3 敏感）\n'
    '    邮箱     → 客户域/客户联系信息（L2 内部）\n\n'
    '4️⃣ 复核员登录，在"人工复核"页面看到3条 AI 映射建议，逐一审核：\n'
    '    → 批准"身份证号 → L4 机密"（正确）\n'
    '    → 批准"手机号 → L3 敏感"（正确）\n'
    '    → 拒绝"邮箱 → L2 内部"（复核员认为应该是 L3，手动调整）\n\n'
    '5️⃣ 批准的映射自动写入系统，字段从"未映射"变为"已映射"，资产目录树自动更新统计数量\n\n'
    '6️⃣ 管理员在报表页面看到：本月新增 200 个字段，AI 映射准确率 85%，人工修正 15%'
)

# ===== 四、当前有哪些功能 =====
doc.add_heading('四、当前版本已实现的功能', level=1)

doc.add_paragraph('这是第一版测试版（MVP），聚焦核心流程：')

features = [
    '用户注册与登录（5个角色，权限分级）',
    '资产目录树管理（多层级、可折叠、拖拽排序）',
    '字段管理（手动录入 + Excel 批量导入导出）',
    '字段与目录的多对多映射（手动 + AI 自动）',
    'AI 智能映射：将未映射字段与目录树发给通义千问大模型，返回分类和分级建议',
    '人工复核工作流：AI 建议需人工批准才生效（异常字段同理）',
    '内置金融数据分类标准：8大分类（客户/交易/产品/财务/运营/风险/合规/系统）+ 子分类',
    '内置 L1-L4 安全分级规则：基于关键字和正则表达式的自动定级引擎',
    '数据打标：为字段添加多维度标签',
    '统计报表：按目录、按敏感级别的分布图表 + Excel 导出',
    '操作日志：记录所有关键操作（查询、筛选）',
    '合规审计页面：敏感字段分布、未映射字段清单、分级合理性检查',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('下个版本会增加：')
run.bold = True
doc.add_paragraph(
    '数据源连接（直连数据库自动扫描字段）、自定义分类标准导入、'
    '批量打标、字段血缘关系图、监管报送自动生成', style='List Bullet'
)

# ===== 五、如何测试 =====
doc.add_heading('五、如何测试', level=1)

doc.add_heading('测试地址', level=2)
doc.add_paragraph(
    '浏览器打开：http://8.130.125.243/dam/\n'
    '（手机浏览器或电脑浏览器都可以，建议用电脑浏览器获得更好体验）'
)

doc.add_heading('测试账号', level=2)

# 账号表格
table = doc.add_table(rows=6, cols=4)
table.style = 'Light Grid Accent 1'
headers = ['角色', '用户名', '密码', '说明']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True

accounts = [
    ['超级管理员', 'admin', 'admin123', '全部权限，用户管理、报表'],
    ['系统管理员', 'system_admin', 'admin123', '分类分级标准管理'],
    ['数据管理员', 'data_admin', 'admin123', '目录管理、AI映射'],
    ['复核员', 'reviewer', 'admin123', '审核AI建议、异常字段'],
    ['录入员', 'data_entry', 'admin123', '录入字段、Excel导入'],
]
for row_idx, row_data in enumerate(accounts):
    for col_idx, val in enumerate(row_data):
        table.rows[row_idx + 1].cells[col_idx].text = val

doc.add_paragraph()

doc.add_heading('推荐测试流程（场景一：体验 AI 分类分级）', level=2)
doc.add_paragraph('建议按以下步骤走一遍，体验"录入 → AI映射 → 人工复核"的完整流程：')

steps = [
    '用"数据管理员"（data_admin）登录 → 点击左侧"资产目录" → 查看已预置的示例目录树（客户域/交易域/风控域 等）',
    '点击"映射管理" → 进入"AI自动映射"标签页 → 点击"执行AI映射"按钮 → 等待几秒，系统会将10个示例字段发给大模型分析',
    'AI 返回映射建议后 → 切换到"AI映射审核"标签页 → 看到每条建议的"映射目标"和"推荐理由"',
    '退出 → 用"复核员"（reviewer）登录 → 点击"人工复核" → 看到AI映射建议和异常字段两个标签页',
    '在"AI映射审核"中逐条审核 → 点击"批准"或"拒绝"（可以填写审核意见）→ 批准的映射自动生效',
    '回到"映射管理" → 查看"已映射"标签 → 刚才批准的映射已经生效',
    '退出 → 用"超级管理员"（admin）登录 → 点击"报表" → 查看按目录、按敏感级别的统计图表',
]
for i, step in enumerate(steps):
    doc.add_paragraph(f'第{i+1}步：{step}')

doc.add_heading('推荐测试流程（场景二：手动录入新字段）', level=2)

steps2 = [
    '用"录入员"（data_entry）登录 → 点击"字段管理" → 点击"新增字段"',
    '填一个测试字段：字段名"银行卡号"，数据类型 VARCHAR，表名 dim_customer，数据库 ods，业务域"客户域"',
    '保存后，系统自动检测：新字段暂未映射 → 标记为异常',
    '在"字段管理"页面可以看到新增的字段，点击"Excel导入"可以批量上传字段',
    '退出 → 用"数据管理员"（data_admin）登录 → "映射管理" → 手动将这个新字段映射到"客户域/客户标识信息"目录下',
    '映射完成后，字段状态从"未映射"变为"已映射"',
]
for i, step in enumerate(steps2):
    doc.add_paragraph(f'第{i+1}步：{step}')

# ===== 六、核心设计理念 =====
doc.add_heading('六、关键设计说明', level=1)

designs = [
    ('AI 辅助，人工决策',
     '平台的 AI 不会直接修改数据分类结果。所有 AI 建议（自动映射、自动分级）都进入"人工复核"队列，'
     '由复核员逐条审批。这既是监管合规的要求（数据分级必须有人的判断），也是对 AI 幻觉的防护。'),
    ('内置行业标准',
     '平台预置了金融行业数据分类的 8 大类别（参考 JR/T 0197-2020《金融数据安全 数据安全分级指南》）'
     '和 L1-L4 四级安全分级规则。系统管理员可以根据机构自身情况修改和扩展这些标准。'),
    ('可追溯、可审计',
     '所有操作（谁在什么时候做了什么）都记录在操作日志中。合规审计页面提供敏感字段分布、'
     '未映射字段清单、分级合理性检查等审计视图，方便应对监管检查。'),
    ('规则引擎 + AI 双引擎',
     '安全分级同时使用两种方式：① 规则引擎（关键字匹配 + 正则表达式）——精确、可解释、可控；'
     '② AI 推理 —— 处理规则覆盖不到的模糊情况。两者结果都进入复核流程。'),
]
for title, desc in designs:
    p = doc.add_paragraph()
    run = p.add_run(f'▸ {title}：')
    run.bold = True
    p.add_run(desc)

# ===== 七、注意事项 =====
doc.add_heading('七、注意事项', level=1)

notes = [
    '这是测试版，界面和功能还会迭代。现在的重点是验证"录入 → AI分类 → 人工复核 → 报表"这个核心流程。',
    '测试数据是预置的示例数据（10个字段、8个目录、2级分类标准），可以随意修改、删除、新增，不影响正式使用。',
    'AI 映射调用的是阿里云通义千问（qwen-plus），响应时间通常在 3-8 秒。如果网络抖动可能稍慢，请耐心等待。',
    '当前版本使用预置示例数据做 AI 映射测试。正式版将支持连接实际数据库，自动扫描表和字段。',
    '如果有任何不好的体验或建议，随时记下来——这些反馈是下个版本优化的方向。',
]
for note in notes:
    doc.add_paragraph(note, style='List Bullet')

# ===== 八、后续计划 =====
doc.add_heading('八、后续计划', level=1)

doc.add_paragraph(
    '当前是第一版（MVP），目标是跑通"AI分类分级 + 人工复核"的核心流程，拿反馈。\n\n'
    '接下来的迭代方向：\n'
    '• 数据源直连——连接 MySQL/PostgreSQL/Oracle 数据库，自动扫描表结构，批量导入字段\n'
    '• 自定义标准导入——支持导入机构自身的分类分级标准（Excel/JSON），与内置标准并行使用\n'
    '• 批量打标——按目录、按敏感级别批量给字段打标签\n'
    '• 字段血缘——展示字段从源表到报表的完整流转路径\n'
    '• 监管报送——根据分类分级结果自动生成监管要求的报送文件\n'
    '• 多租户——不同部门/分公司独立管理各自的数据资产目录\n\n'
    '如果有其他想法，随时沟通。'
)

# ===== 保存 =====
import os
docs_dir = r'D:\AI projects\data-asset-management-platform\docs'
os.makedirs(docs_dir, exist_ok=True)
output_path = os.path.join(docs_dir, '数据资产管理平台-测试版操作指南.docx')
doc.save(output_path)
print(f'文档已生成：{output_path}')
